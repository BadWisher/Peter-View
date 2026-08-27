"""Парсинг входных форматов в единую блочную структуру.

На входе — docx, html, markdown, txt или URL. На выходе — Document из блоков,
где у каждого блока хранятся отдельно сырой текст с разметкой (raw) и чистый
текст (plain), плюс метаданные стиля/уровня.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import httpx
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document as DocxDocument
from docx.document import Document as _DocxDocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt

from ..extractors import assert_docx_safe, extract_html, normalize_spaces
from ..net_guard import safe_get

BLOCK_HTML_TAGS = (
    "h1", "h2", "h3", "h4", "h5", "h6",
    "p", "li", "blockquote", "pre", "tr",
    "figcaption", "dd", "dt", "caption",
)

STRUCTURED_MAX_CHARS = 16_000
STRUCTURED_BLOCK_MAX_CHARS = 2_000
_MAX_LINKS_PER_BLOCK = 20
_MAX_CELLS_PER_ROW = 40

_MD_STRIP_PATTERNS = [
    (re.compile(r"`([^`]+)`"), r"\1"),
    (re.compile(r"!\[([^\]]*)\]\([^)]+\)"), r"\1"),
    (re.compile(r"\[([^\]]+)\]\([^)]+\)"), r"\1"),
    (re.compile(r"^#{1,6}\s+", re.MULTILINE), ""),
    (re.compile(r"\*\*(.+?)\*\*"), r"\1"),
    (re.compile(r"(?<!\*)\*(?!\*)(.+?)\*"), r"\1"),
    (re.compile(r"^>\s?", re.MULTILINE), ""),
    (re.compile(r"^[\s]*[-*+]\s+", re.MULTILINE), ""),
    (re.compile(r"^[\s]*\d+\.\s+", re.MULTILINE), ""),
]


@dataclass
class Block:
    index: int
    raw: str
    plain: str
    metadata: dict = field(default_factory=dict)

    def structured(self, max_chars: int = STRUCTURED_BLOCK_MAX_CHARS) -> str:
        """Каноническое ограниченное JSON-представление блока для LLM."""
        return _bounded_json(_block_payload(self), max_chars)


@dataclass
class Document:
    blocks: list[Block]
    source: str = ""

    def full_plain(self) -> str:
        return "\n".join(b.plain for b in self.blocks if b.plain)

    def full_raw(self) -> str:
        return "\n".join(b.raw for b in self.blocks if b.raw)

    def full_structured(self, max_chars: int = STRUCTURED_MAX_CHARS) -> str:
        """Канонический JSON документа с жёстким ограничением размера."""
        max_chars = max(2, max_chars)
        payload: dict = {"source": self.source, "blocks": []}
        for block in self.blocks:
            candidate = payload["blocks"] + [_block_payload(block)]
            encoded = json.dumps(
                {"source": self.source, "blocks": candidate},
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )
            if len(encoded) > max_chars:
                payload["truncated"] = True
                payload["remaining_blocks"] = len(self.blocks) - len(payload["blocks"])
                break
            payload["blocks"] = candidate
        encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
        if len(encoded) <= max_chars:
            return encoded
        # Even metadata about truncation must respect the caller's bound.
        return _bounded_json(
            {"source": self.source, "blocks": [], "truncated": True,
             "remaining_blocks": len(self.blocks)},
            max_chars,
        )


def _trim(value: object, limit: int = 1_000) -> object:
    if isinstance(value, str):
        return value if len(value) <= limit else value[: max(0, limit - 1)] + "…"
    if isinstance(value, list):
        return [_trim(item, limit) for item in value[:_MAX_CELLS_PER_ROW]]
    if isinstance(value, dict):
        return {str(key): _trim(item, limit) for key, item in value.items()}
    return value


def _block_payload(block: Block) -> dict:
    metadata = _trim(block.metadata)
    if isinstance(metadata, dict) and isinstance(metadata.get("links"), list):
        metadata["links"] = metadata["links"][:_MAX_LINKS_PER_BLOCK]
    return {
        "index": block.index,
        "type": block.metadata.get("type", "paragraph"),
        "plain": _trim(block.plain),
        "raw": _trim(block.raw),
        "metadata": metadata,
    }


def _bounded_json(payload: dict, max_chars: int) -> str:
    max_chars = max(2, max_chars)
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    if len(encoded) <= max_chars:
        return encoded
    fallback = json.dumps({"truncated": True}, separators=(",", ":"))
    if len(fallback) <= max_chars:
        return fallback
    return "{}"


def _links_from_tag(node: Tag) -> list[dict]:
    return [
        {"text": normalize_spaces(link.get_text(" ", strip=True)), "target": link.get("href", "")}
        for link in node.find_all("a", href=True)
    ][:_MAX_LINKS_PER_BLOCK]


def _tag_format_flags(node: Tag | None) -> dict[str, bool]:
    flags = {"bold": False, "italic": False, "underline": False}
    current = node
    while current is not None:
        name = getattr(current, "name", "") or ""
        if name in {"strong", "b"}:
            flags["bold"] = True
        elif name in {"em", "i"}:
            flags["italic"] = True
        elif name == "u":
            flags["underline"] = True
        current = getattr(current, "parent", None)
    return flags


def _remap_format_spans(source: str, target: str, spans: list[dict]) -> list[dict]:
    if not spans:
        return []
    if source == target:
        return spans
    remapped: list[dict] = []
    cursor = 0
    for span in spans:
        snippet = normalize_spaces(source[span["start"]:span["end"]])
        if not snippet:
            continue
        index = target.find(snippet, cursor)
        if index < 0:
            index = target.find(snippet)
        if index < 0:
            continue
        remapped.append({
            "start": index,
            "end": index + len(snippet),
            "bold": bool(span.get("bold")),
            "italic": bool(span.get("italic")),
            "underline": bool(span.get("underline")),
        })
        cursor = index + len(snippet)
    return remapped


def _formatting_from_html(node: Tag, plain: str, list_owner: Tag | None = None) -> list[dict]:
    parts: list[str] = []
    spans: list[dict] = []
    offset = 0
    for descendant in node.descendants:
        if not isinstance(descendant, NavigableString):
            continue
        if list_owner is not None and descendant.find_parent(("ul", "ol")) is not list_owner:
            continue
        text = str(descendant)
        if not text:
            continue
        flags = _tag_format_flags(descendant.parent if isinstance(descendant.parent, Tag) else None)
        if any(flags.values()):
            spans.append({"start": offset, "end": offset + len(text), **flags})
        parts.append(text)
        offset += len(text)
    return _remap_format_spans("".join(parts), plain, spans)


def _md_inline_formatting(token, target_plain: str) -> list[dict]:
    children = getattr(token, "children", None) or []
    parts: list[str] = []
    spans: list[dict] = []
    offset = 0
    bold = italic = False
    for child in children:
        kind = child.type
        if kind == "strong_open":
            bold = True
        elif kind == "strong_close":
            bold = False
        elif kind == "em_open":
            italic = True
        elif kind == "em_close":
            italic = False
        elif kind in {"text", "code_inline"}:
            text = child.content or ""
            if text and (bold or italic):
                spans.append({
                    "start": offset,
                    "end": offset + len(text),
                    "bold": bold,
                    "italic": italic,
                    "underline": False,
                })
            parts.append(text)
            offset += len(text)
        elif kind in {"softbreak", "hardbreak"}:
            parts.append("\n")
            offset += 1
    return _remap_format_spans("".join(parts), target_plain, spans)


def _attach_list_introductions(blocks: list[Block]) -> None:
    """Связывает список с непосредственно предшествующим вводным блоком."""
    group = 0
    pos = 0
    while pos < len(blocks):
        block = blocks[pos]
        if block.metadata.get("type") != "list_item" or block.metadata.get("list_depth", 0) != 0:
            pos += 1
            continue
        group += 1
        start = pos
        while pos < len(blocks) and blocks[pos].metadata.get("type") == "list_item":
            blocks[pos].metadata.setdefault("list_group", group)
            pos += 1
        if start and blocks[start - 1].metadata.get("type") in {"paragraph", "blockquote"}:
            intro = blocks[start - 1]
            intro.metadata["introduces_list"] = blocks[start].index
            for item in blocks[start:pos]:
                item.metadata["list_intro_index"] = intro.index


def _md_to_plain(raw: str) -> str:
    text = raw
    for pattern, repl in _MD_STRIP_PATTERNS:
        text = pattern.sub(repl, text)
    return text.strip()


def parse_markdown(text: str, source: str = "") -> Document:
    md = MarkdownIt().enable("table")
    lines = text.split("\n")
    blocks: list[Block] = []
    tokens = md.parse(text)
    list_stack: list[dict] = []
    item_stack: list[dict] = []
    skip_inline: set[int] = set()

    table_index = -1
    table_start = 0
    row_index = 0
    cells: list[str] = []
    row_links: list[dict] = []
    in_table = False
    for token in tokens:
        if token.type == "table_open":
            table_index += 1
            table_start = token.map[0] if token.map else 0
            row_index = 0
            in_table = True
        elif token.type == "table_close":
            in_table = False
        elif token.type == "tr_open" and in_table:
            cells = []
            row_links = []
        elif token.type == "inline" and in_table:
            cells.append(_md_to_plain(token.content))
            row_links.extend(
                {"text": label, "target": target}
                for label, target in re.findall(
                    r"\[([^\]]+)\]\(([^)\s]+)(?:\s+[^)]*)?\)", token.content
                )
            )
        elif token.type == "tr_close" and in_table:
            # Markdown's separator consumes one source line after the header.
            source_line = table_start + row_index + (1 if row_index else 0)
            metadata: dict = {
                "type": "table_row",
                "table_index": table_index,
                "row_index": row_index,
                "cells": cells[:_MAX_CELLS_PER_ROW],
                "_source_line": source_line,
            }
            if row_links:
                metadata["links"] = row_links[:_MAX_LINKS_PER_BLOCK]
            raw = lines[source_line].strip() if source_line < len(lines) else " | ".join(cells)
            blocks.append(Block(len(blocks), raw, " | ".join(cells), metadata))
            row_index += 1

    for token_pos, token in enumerate(tokens):
        if token.type in {"bullet_list_open", "ordered_list_open"}:
            list_stack.append({
                "kind": "ordered" if token.type == "ordered_list_open" else "bullet",
                "next_index": int(token.attrGet("start") or 1),
            })
        elif token.type in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
        elif token.type == "list_item_open":
            current = list_stack[-1]
            item_stack.append({
                "kind": current["kind"],
                "depth": len(list_stack) - 1,
                "index": current["next_index"],
                "map": token.map,
                "emitted": False,
            })
            current["next_index"] += 1
        elif token.type == "list_item_close":
            if item_stack:
                item_stack.pop()
        elif token.type == "inline" and item_stack and not item_stack[-1]["emitted"]:
            item = item_stack[-1]
            start, end = token.map or item["map"] or (0, 0)
            raw = "\n".join(lines[start:end]).strip()
            metadata = {
                "type": "list_item",
                "list_kind": item["kind"],
                "list_depth": item["depth"],
                "list_index": item["index"],
                "_source_line": start,
            }
            links = [
                {"text": child.content, "target": child.attrGet("href") or ""}
                for child in token.children or []
                if child.type == "link_open"
            ]
            # markdown-it stores link text in following text children, so recover it
            # from the source for a stable useful label.
            source_links = re.findall(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+[^)]*)?\)", token.content)
            if source_links:
                links = [{"text": label, "target": target} for label, target in source_links]
            if links:
                metadata["links"] = links[:_MAX_LINKS_PER_BLOCK]
            plain = _md_to_plain(token.content)
            formatting = _md_inline_formatting(token, plain)
            if formatting:
                metadata["formatting"] = formatting
            blocks.append(Block(len(blocks), raw, plain, metadata))
            item["emitted"] = True
            skip_inline.add(token_pos)

    # Emit non-list block tokens separately. Their source maps preserve raw Markdown.
    for token_pos, token in enumerate(tokens):
        if token_pos in skip_inline or not token.map:
            continue
        if token.type not in {
            "heading_open", "paragraph_open", "blockquote_open",
            "fence", "code_block", "hr",
        }:
            continue
        if token.level and token.type not in {"fence", "code_block"}:
            # Nested list paragraphs have already been represented above.
            continue
        start, end = token.map
        raw = "\n".join(lines[start:end]).strip()
        if not raw:
            continue
        metadata: dict = {"type": token.type.replace("_open", "")}
        if token.type == "heading_open":
            metadata = {"type": "heading", "level": int(token.tag[1:])}
        elif token.type == "hr":
            metadata = {"type": "horizontal_rule"}
        metadata["_source_line"] = start
        inline = tokens[token_pos + 1] if token_pos + 1 < len(tokens) else None
        plain = _md_to_plain(raw)
        if inline and inline.type == "inline":
            links = [
                {"text": label, "target": target}
                for label, target in re.findall(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+[^)]*)?\)", inline.content)
            ]
            if links:
                metadata["links"] = links[:_MAX_LINKS_PER_BLOCK]
            formatting = _md_inline_formatting(inline, plain)
            if formatting:
                metadata["formatting"] = formatting
        blocks.append(Block(len(blocks), raw, plain, metadata))

    # Token passes above are optimized for metadata; restore source order from maps.
    blocks.sort(key=lambda block: (block.metadata.pop("_source_line", len(lines)), block.index))
    for index, block in enumerate(blocks):
        block.index = index
    _attach_list_introductions(blocks)

    return Document(blocks=blocks, source=source)
def parse_html(html: str, source: str = "") -> Document:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup.find_all(("script", "style", "noscript", "head", "nav", "footer", "header")):
        tag.decompose()

    body = soup.find("body") or soup
    blocks: list[Block] = []
    idx = 0

    for node in body.find_all(BLOCK_HTML_TAGS):
        owner = None
        if node.name == "p" and node.find_parent(("li", "td", "th")):
            continue
        if node.name in {"blockquote", "pre"} and node.find_parent(("blockquote", "pre")):
            continue
        if node.name == "li":
            owner = node.find_parent(("ul", "ol"))
            parts = [
                descendant
                for descendant in node.descendants
                if isinstance(descendant, NavigableString)
                and descendant.find_parent(("ul", "ol")) is owner
            ]
            plain = normalize_spaces(" ".join(str(part) for part in parts))
        else:
            plain = normalize_spaces(node.get_text(separator=" ", strip=True))
        if not plain:
            continue
        raw = node.decode().strip()
        metadata: dict = {"type": node.name}
        if node.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            metadata = {"type": "heading", "level": int(node.name[1])}
        elif node.name == "li":
            metadata = {
                "type": "list_item",
                "list_kind": "ordered" if owner and owner.name == "ol" else "bullet",
                "list_depth": max(0, len(node.find_parents(("ul", "ol"))) - 1),
                "list_index": (
                    int(node.get("value"))
                    if node.get("value", "").lstrip("-").isdigit()
                    else int(owner.get("start", 1) if owner else 1)
                    + len(node.find_previous_siblings("li"))
                ),
            }
            raw = "".join(
                str(child) for child in node.contents
                if not isinstance(child, Tag) or child.name not in {"ul", "ol"}
            ).strip()
        elif node.name == "tr":
            table = node.find_parent("table")
            cells = [normalize_spaces(cell.get_text(" ", strip=True)) for cell in node.find_all(("td", "th"), recursive=False)]
            metadata = {
                "type": "table_row",
                "table_index": len(table.find_previous_siblings("table")) if table else 0,
                "row_index": len(node.find_previous_siblings("tr")),
                "cells": cells[:_MAX_CELLS_PER_ROW],
            }
            plain = " | ".join(cells)
            raw = node.decode().strip()
        links = _links_from_tag(node)
        if links:
            metadata["links"] = links
        formatting = _formatting_from_html(
            node, plain, list_owner=owner if node.name == "li" else None
        )
        if formatting:
            metadata["formatting"] = formatting
        blocks.append(Block(index=idx, raw=raw, plain=plain, metadata=metadata))
        idx += 1

    if not blocks:
        text = extract_html(html)
        blocks = [
            Block(index=i, raw=line, plain=line, metadata={"type": "paragraph"})
            for i, line in enumerate(filter(None, (ln.strip() for ln in text.split("\n"))))
        ]

    _attach_list_introductions(blocks)
    return Document(blocks=blocks, source=source)


def _iter_docx_blocks(parent: _DocxDocumentType) -> Iterable[Paragraph | Table]:
    for child in parent.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _docx_text_and_links(paragraph: Paragraph) -> tuple[str, list[dict], list[dict]]:
    parts: list[str] = []
    links: list[dict] = []
    formatting: list[dict] = []
    offset = 0
    for child in paragraph._p.iterchildren():
        kind = child.tag.rsplit("}", 1)[-1]
        text = "".join(node.text or "" for node in child.iter(qn("w:t")))
        parts.append(text)
        for run in child.iter(qn("w:r")):
            run_text = "".join(node.text or "" for node in run.iter(qn("w:t")))
            if not run_text:
                continue
            properties = run.find(qn("w:rPr"))
            if properties is not None:
                flags = {
                    "bold": properties.find(qn("w:b")) is not None,
                    "italic": properties.find(qn("w:i")) is not None,
                    "underline": properties.find(qn("w:u")) is not None,
                }
                if any(flags.values()):
                    formatting.append({
                        "start": offset,
                        "end": offset + len(run_text),
                        **flags,
                    })
            offset += len(run_text)
        if kind == "hyperlink":
            rel_id = child.get(qn("r:id"))
            anchor = child.get(qn("w:anchor"))
            target = f"#{anchor}" if anchor else ""
            if rel_id and rel_id in paragraph.part.rels:
                target = paragraph.part.rels[rel_id].target_ref
            links.append({"text": text, "target": target})
    return "".join(parts).strip(), links[:_MAX_LINKS_PER_BLOCK], formatting


def _docx_list_metadata(
    paragraph: Paragraph,
    counters: dict[tuple[object, int], int],
) -> dict | None:
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None else None
    style = paragraph.style.name if paragraph.style else "Normal"
    if num_pr is None or num_pr.numId is None:
        if style.lower().startswith("list"):
            depth_match = re.search(r"(\d+)$", style)
            depth = max(0, int(depth_match.group(1)) - 1) if depth_match else 0
            key = (style.lower(), depth)
            counters[key] = counters.get(key, 0) + 1
            return {
                "type": "list_item",
                "list_kind": "ordered" if "number" in style.lower() else "bullet",
                "list_depth": depth,
                "list_index": counters[key],
            }
        return None
    num_id = int(num_pr.numId.val)
    depth = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
    formats = paragraph.part.numbering_part.element.xpath(
        f"./w:num[@w:numId='{num_id}']/w:abstractNumId"
    )
    kind = "ordered"
    if formats:
        abstract_id = formats[0].get(qn("w:val"))
        num_formats = paragraph.part.numbering_part.element.xpath(
            f"./w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='{depth}']/w:numFmt"
        )
        if num_formats and num_formats[0].get(qn("w:val")) == "bullet":
            kind = "bullet"
    key = (num_id, depth)
    counters[key] = counters.get(key, 0) + 1
    return {
        "type": "list_item",
        "list_kind": kind,
        "list_depth": depth,
        "list_index": counters[key],
        "list_id": num_id,
    }


def parse_docx(content: bytes, source: str = "") -> Document:
    assert_docx_safe(content)
    doc = DocxDocument(io.BytesIO(content))
    blocks: list[Block] = []
    list_counters: dict[tuple[object, int], int] = {}
    table_index = 0

    for item in _iter_docx_blocks(doc):
        if isinstance(item, Paragraph):
            text, links, formatting = _docx_text_and_links(item)
            if not text:
                continue
            style = item.style.name if item.style else "Normal"
            metadata: dict = {"type": "paragraph", "style": style}
            list_metadata = _docx_list_metadata(item, list_counters)
            if list_metadata:
                metadata.update(list_metadata)
            elif style.lower().startswith("heading"):
                level = "".join(ch for ch in style if ch.isdigit())
                metadata.update({"type": "heading", "level": int(level) if level else 0})
            if links:
                metadata["links"] = links
            if formatting:
                metadata["formatting"] = formatting
            blocks.append(Block(index=len(blocks), raw=text, plain=text, metadata=metadata))
            continue

        for row_index, row in enumerate(item.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            line = " | ".join(cells)
            links = []
            formatting = []
            for cell_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    _, paragraph_links, paragraph_formatting = _docx_text_and_links(paragraph)
                    links.extend({**link, "cell_index": cell_index} for link in paragraph_links)
                    formatting.extend(
                        {**span, "cell_index": cell_index}
                        for span in paragraph_formatting
                    )
            metadata = {
                "type": "table_row",
                "table_index": table_index,
                "row_index": row_index,
                "cells": cells[:_MAX_CELLS_PER_ROW],
            }
            if links:
                metadata["links"] = links[:_MAX_LINKS_PER_BLOCK]
            if formatting:
                metadata["formatting"] = formatting
            blocks.append(Block(index=len(blocks), raw=line, plain=line, metadata=metadata))
        table_index += 1

    _attach_list_introductions(blocks)
    return Document(blocks=blocks, source=source)


def parse_txt(text: str, source: str = "") -> Document:
    blocks: list[Block] = []
    paragraph: list[str] = []
    list_counters: dict[tuple[str, int], int] = {}
    list_re = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.+)$")

    def flush_paragraph() -> None:
        if not paragraph:
            return
        value = "\n".join(paragraph).strip()
        paragraph.clear()
        if value:
            blocks.append(Block(
                index=len(blocks),
                raw=value,
                plain=value,
                metadata={"type": "paragraph"},
            ))

    for line in text.splitlines():
        if not line.strip():
            flush_paragraph()
            continue
        match = list_re.match(line)
        if not match:
            paragraph.append(line.strip())
            continue
        flush_paragraph()
        indent, marker, value = match.groups()
        depth = len(indent.expandtabs(2)) // 2
        kind = "bullet" if marker in {"-", "*", "+"} else "ordered"
        if kind == "ordered":
            list_index = int(re.match(r"\d+", marker).group())
        else:
            key = (kind, depth)
            list_counters[key] = list_counters.get(key, 0) + 1
            list_index = list_counters[key]
        blocks.append(Block(
            index=len(blocks),
            raw=line.strip(),
            plain=value.strip(),
            metadata={
                "type": "list_item",
                "list_kind": kind,
                "list_depth": depth,
                "list_index": list_index,
            },
        ))
    flush_paragraph()
    _attach_list_introductions(blocks)
    return Document(blocks=blocks, source=source)


def parse_file(content: bytes, filename: str) -> Document:
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return parse_docx(content, source=filename)
    if ext in (".html", ".htm"):
        return parse_html(content.decode("utf-8", errors="replace"), source=filename)
    if ext == ".md":
        return parse_markdown(content.decode("utf-8", errors="replace"), source=filename)
    if ext == ".txt":
        return parse_txt(content.decode("utf-8", errors="replace"), source=filename)
    raise ValueError(f"Неподдерживаемый формат файла: {ext}")


async def parse_url(url: str) -> Document:
    async with httpx.AsyncClient(
        timeout=30.0,
        headers={"User-Agent": "Mozilla/5.0 (compatible; Proofreader/1.0)"},
    ) as client:
        resp = await safe_get(client, url)
        resp.raise_for_status()
    return parse_html(resp.text, source=url)
