import io
import json
import unittest

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

from app.llm.documents import parse_docx, parse_html, parse_markdown, parse_txt


def _docx_bytes(document: DocxDocument) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _add_hyperlink(paragraph, text: str, target: str) -> None:
    relation_id = paragraph.part.relate_to(
        target, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class DocumentParsingTests(unittest.TestCase):
    def test_markdown_captures_heading_nested_lists_links_and_intro(self):
        document = parse_markdown(
            "# Setup\n\nChoose an option:\n\n"
            "1. First\n2. [Second](https://example.test)\n   - Nested\n\nDone."
        )

        self.assertEqual(
            [block.metadata["type"] for block in document.blocks],
            ["heading", "paragraph", "list_item", "list_item", "list_item", "paragraph"],
        )
        items = document.blocks[2:5]
        self.assertEqual([item.metadata["list_depth"] for item in items], [0, 0, 1])
        self.assertEqual(items[1].plain, "Second")
        self.assertEqual(items[1].metadata["links"][0]["target"], "https://example.test")
        self.assertEqual(items[0].metadata["list_intro_index"], 1)
        self.assertEqual(document.blocks[1].metadata["introduces_list"], 2)

    def test_markdown_captures_bold_and_italic_spans(self):
        document = parse_markdown("A **bold** and *italic* word.")
        paragraph = document.blocks[0]
        self.assertEqual(paragraph.plain, "A bold and italic word.")
        self.assertEqual(paragraph.metadata["formatting"], [
            {"start": 2, "end": 6, "bold": True, "italic": False, "underline": False},
            {"start": 11, "end": 17, "bold": False, "italic": True, "underline": False},
        ])

    def test_html_captures_bold_and_italic_spans(self):
        document = parse_html("<p>A <strong>bold</strong> and <em>italic</em> word.</p>")
        paragraph = document.blocks[0]
        self.assertEqual(paragraph.plain, "A bold and italic word.")
        self.assertEqual(paragraph.metadata["formatting"], [
            {"start": 2, "end": 6, "bold": True, "italic": False, "underline": False},
            {"start": 11, "end": 17, "bold": False, "italic": True, "underline": False},
        ])

    def test_html_captures_lists_links_and_table_positions(self):
        document = parse_html(
            "<h2>Options</h2><p>Use one:</p><ul>"
            "<li>Alpha</li><li><a href='/beta'>Beta</a><ol><li>Nested</li></ol></li>"
            "</ul><table><tr><th>Name</th></tr><tr><td>Value</td></tr></table>"
        )

        self.assertEqual([block.plain for block in document.blocks[2:5]], ["Alpha", "Beta", "Nested"])
        self.assertEqual(document.blocks[4].metadata["list_depth"], 1)
        self.assertEqual(document.blocks[3].metadata["links"], [{"text": "Beta", "target": "/beta"}])
        rows = [block for block in document.blocks if block.metadata["type"] == "table_row"]
        self.assertEqual(
            [(row.metadata["table_index"], row.metadata["row_index"]) for row in rows],
            [(0, 0), (0, 1)],
        )

    def test_markdown_captures_table_rows_in_source_order(self):
        document = parse_markdown(
            "Before\n\n| Name | Target |\n| --- | --- |\n"
            "| Item | [Page](/page) |\n\nAfter"
        )

        self.assertEqual(
            [block.metadata["type"] for block in document.blocks],
            ["paragraph", "table_row", "table_row", "paragraph"],
        )
        rows = document.blocks[1:3]
        self.assertEqual([row.metadata["row_index"] for row in rows], [0, 1])
        self.assertEqual(rows[1].metadata["links"], [{"text": "Page", "target": "/page"}])
        self.assertEqual(rows[1].plain, "Item | Page")

    def test_docx_preserves_paragraph_table_order_and_hyperlinks(self):
        source = DocxDocument()
        source.add_heading("Title", level=1)
        paragraph = source.add_paragraph("See ")
        _add_hyperlink(paragraph, "reference", "https://example.test/reference")
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        source.add_paragraph("After table")

        document = parse_docx(_docx_bytes(source))

        self.assertEqual([block.plain for block in document.blocks], [
            "Title", "See reference", "A | B", "After table"
        ])
        self.assertEqual(document.blocks[2].metadata["table_index"], 0)
        self.assertEqual(
            document.blocks[1].metadata["links"],
            [{"text": "reference", "target": "https://example.test/reference"}],
        )

    def test_docx_captures_list_kind_depth_index_and_introduction(self):
        source = DocxDocument()
        source.add_paragraph("Steps:")
        source.add_paragraph("First", style="List Number")
        source.add_paragraph("Second", style="List Number")

        document = parse_docx(_docx_bytes(source))

        first, second = document.blocks[1:]
        self.assertEqual(first.metadata["list_kind"], "ordered")
        self.assertEqual(first.metadata["list_depth"], 0)
        self.assertEqual([first.metadata["list_index"], second.metadata["list_index"]], [1, 2])
        self.assertEqual(first.metadata["list_intro_index"], 0)

    def test_docx_captures_basic_inline_formatting_spans(self):
        source = DocxDocument()
        paragraph = source.add_paragraph("Normal ")
        run = paragraph.add_run("bold")
        run.bold = True
        italic = paragraph.add_run(" italic")
        italic.italic = True

        document = parse_docx(_docx_bytes(source))

        self.assertEqual(document.blocks[0].plain, "Normal bold italic")
        self.assertEqual(document.blocks[0].metadata["formatting"], [
            {"start": 7, "end": 11, "bold": True, "italic": False, "underline": False},
            {"start": 11, "end": 18, "bold": False, "italic": True, "underline": False},
        ])

    def test_structured_serialization_is_canonical_bounded_and_compatible(self):
        document = parse_txt("First paragraph.\n\nSecond paragraph.", source="sample.txt")

        self.assertEqual(document.full_plain(), "First paragraph.\nSecond paragraph.")
        self.assertEqual(document.full_raw(), document.full_plain())
        structured = document.full_structured(max_chars=1_000)
        payload = json.loads(structured)
        self.assertEqual(payload["source"], "sample.txt")
        self.assertEqual(payload["blocks"][0]["type"], "paragraph")
        self.assertEqual(structured, document.full_structured(max_chars=1_000))
        self.assertLessEqual(len(document.full_structured(max_chars=40)), 40)


if __name__ == "__main__":
    unittest.main()
